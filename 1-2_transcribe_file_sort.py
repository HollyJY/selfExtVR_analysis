# %% [markdown]
# put all transcribed text into a single file, sorted by folder name and file name

# * - [-] ['P02', 'P09', 'P25'] A <-> switch?

# %%
import os
from pathlib import Path

from pathlib import Path
from docx import Document
from docx.shared import Pt

# %%
src_folder_base = Path("data_raw/interviews/")
dst_base = Path("data_intermediate/")

folders_pp = sorted([f for f in src_folder_base.iterdir() if f.is_dir()])

# %%
# copy all clean.txt files to data_intermediate folder
for folder in folders_pp:
    src = folder
    pp = folder.name
    dst_folder = dst_base / folder.name
    for root, dirs, files in os.walk(src):
        for file in files:
            if file.endswith("clean.txt"):
                src_file = Path(root) / file
                dst_file = dst_folder / file
                dst_file.parent.mkdir(parents=True, exist_ok=True)
                os.rename(src_file, dst_file)
                print(f"Moved {src_file} to {dst_file}")
            
# %% 
# file counts

src_base = Path("data_intermediate/")

for index, folder in enumerate(folders_pp):
    folder  = src_base / folder.name
    txt_files = sorted(folder.glob("*.txt"))
    # print(folder.name, txt_files)
    if txt_files:
        txt_file = txt_files[0]
        print(f"{folder.name}: {len(txt_files)} txt files, first file: {txt_file.name}")
    else:
        print(f"{folder.name}: No txt files found.")


# %%
# all txt into 1 file

src_base = Path("data_intermediate/")
dst_base = Path("data_analysis/")
output_file = "data_analysis/interview_all_raw.docx"

document = Document()

# Optional: set default font size
style = document.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(12)

folders_pp = sorted([f for f in src_folder_base.iterdir() if f.is_dir()])

for index, folder in enumerate(folders_pp):
    folder  = src_base / folder.name
    txt_files = sorted(folder.glob("*.txt"))
    # print(folder.name, txt_files)
    if txt_files:
        txt_file = txt_files[0]
        text = txt_file.read_text(encoding="utf-8")

    # File name as heading
    document.add_heading(folder.name, level=1)

    # Preserve paragraphs
    for paragraph in text.split("\n"):
        document.add_paragraph(paragraph)

    # Each txt file starts on a new page
    if index < len(folders_pp) - 1:
        document.add_page_break()

document.save(output_file)

print(f"Created: {output_file}")
# %%
